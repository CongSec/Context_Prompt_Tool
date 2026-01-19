# ai.py
# -*- coding: utf-8 -*-
import csv
import io
import json
import os
import re
import sqlite3
import threading
import time
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
from uuid import uuid4
from dataclasses import dataclass, field, asdict
from typing import Any, Dict, List, Optional, Tuple, Union

import requests
from PyQt5 import QtCore, QtGui, QtWidgets

# 可选依赖，用于支持 office 文档
# 用户需自行安装：pip install python-docx openpyxl xlrd
try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import openpyxl  # xlsx
except Exception:
    openpyxl = None

try:
    import xlrd  # xls
except Exception:
    xlrd = None

# 增强的 .doc 文件支持
try:
    from win32com.client import Dispatch
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False

try:
    import subprocess
    import sys
    SUBPROCESS_AVAILABLE = True
except ImportError:
    SUBPROCESS_AVAILABLE = False

# 新增：尝试导入其他可选的 .doc 处理库
try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False

try:
    from antiword import antiword
    ANTIWORD_AVAILABLE = True
except ImportError:
    ANTIWORD_AVAILABLE = False

APP_NAME = "上下文提示词工具 by CongSec"
DEFAULT_CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")

ID_REGEX = re.compile(r"\b\d{14}-[a-z0-9]{6,}\b", re.IGNORECASE)
FILE_SIZE_LIMIT = 20 * 1024 * 1024  # 20MB per file
MAX_CONCURRENT_FILES = 5  # 最大并发文件处理数

ALLOWED_TEXT_EXTS = {
    ".txt", ".md", ".py", ".java", ".c", ".h", ".cpp", ".hpp", ".js", ".ts",
    ".json", ".yml", ".yaml", ".xml", ".html", ".css", ".go", ".rs", ".sh",
    ".bat", ".ps1", ".sql", ".ini", ".cfg", ".toml", ".log", ".csv", ".tsv",
    ".doc", ".docx", ".xls", ".xlsx"
}


def now_label() -> str:
    return QtCore.QDateTime.currentDateTime().toString("yyyy-MM-dd hh:mm:ss")


def safe_filename(name: str) -> str:
    name = re.sub(r"[^\w\-. ]+", "_", name.strip())
    return name if name else "untitled"


def normalize_id_list(text: str) -> List[str]:
    if not text:
        return []
    found = ID_REGEX.findall(text)
    return list(dict.fromkeys(found))  # unique preserve order


def read_text_files_by_paths(file_paths: List[str]) -> Dict[str, Tuple[str, int]]:
    """优化版本：使用线程池并发处理文件读取"""
    results: Dict[str, Tuple[str, int]] = {}
    
    def read_single_file(fp: str) -> Tuple[str, int, str]:
        try:
            size = os.path.getsize(fp)
            if size > FILE_SIZE_LIMIT:
                return fp, size, f"[文件过大，未读取: {os.path.basename(fp)} ({size} bytes)]"
            
            content = read_text_from_any_file(fp)
            return fp, size, content
        except Exception as e:
            return fp, 0, f"[读取失败: {os.path.basename(fp)}]\n{str(e)}"
    
    # 限制并发数量，避免同时打开太多文件
    with ThreadPoolExecutor(max_workers=MAX_CONCURRENT_FILES) as executor:
        future_to_file = {executor.submit(read_single_file, fp): fp for fp in file_paths}
        
        for future in as_completed(future_to_file):
            fp, size, content = future.result()
            results[fp] = (content, size)
    
    return results


def read_doc_file_with_win32com(path: str) -> str:
    """使用 win32com 读取 .doc 文件（Windows 系统）"""
    try:
        word = Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(path))
        content = doc.Content.Text
        doc.Close()
        word.Quit()
        return content.strip()
    except Exception as e:
        return f"[win32com 读取失败: {str(e)}]"


def read_doc_file_with_antiword(path: str) -> str:
    """使用 antiword 命令行工具读取 .doc 文件"""
    try:
        result = subprocess.run(
            ['antiword', path], 
            capture_output=True, 
            text=True, 
            timeout=30,
            encoding='utf-8',
            errors='ignore'
        )
        if result.returncode == 0:
            return result.stdout
        else:
            return f"[antiword 转换失败: {result.stderr}]"
    except FileNotFoundError:
        return "[未找到 antiword 命令，请安装 antiword]"
    except subprocess.TimeoutExpired:
        return "[antiword 转换超时]"
    except Exception as e:
        return f"[antiword 转换异常: {str(e)}]"


def read_doc_file_with_catdoc(path: str) -> str:
    """使用 catdoc 命令行工具读取 .doc 文件"""
    try:
        result = subprocess.run(
            ['catdoc', path], 
            capture_output=True, 
            text=True, 
            timeout=30,
            encoding='utf-8',
            errors='ignore'
        )
        if result.returncode == 0:
            return result.stdout
        else:
            return f"[catdoc 转换失败: {result.stderr}]"
    except FileNotFoundError:
        return "[未找到 catdoc 命令，请安装 catdoc]"
    except subprocess.TimeoutExpired:
        return "[catdoc 转换超时]"
    except Exception as e:
        return f"[catdoc 转换异常: {str(e)}]"


def read_doc_file_with_unoconv(path: str) -> str:
    """使用 unoconv 命令行工具将 .doc 转换为文本"""
    try:
        # 检查 unoconv 是否可用
        result = subprocess.run(['unoconv', '--version'], 
                              capture_output=True, text=True, timeout=10)
        if result.returncode != 0:
            return "[unoconv 不可用，请安装 LibreOffice 和 unoconv]"
        
        # 创建临时文件路径
        import tempfile
        temp_dir = tempfile.mkdtemp()
        temp_txt = os.path.join(temp_dir, 'output.txt')
        
        # 转换文件
        cmd = ['unoconv', '-f', 'txt', '-o', temp_dir, path]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0 and os.path.exists(temp_txt):
            with open(temp_txt, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            # 清理临时文件
            import shutil
            shutil.rmtree(temp_dir)
            return content
        else:
            if os.path.exists(temp_txt):
                import shutil
                shutil.rmtree(temp_dir)
            return f"[unoconv 转换失败: {result.stderr}]"
    except subprocess.TimeoutExpired:
        return "[unoconv 转换超时]"
    except FileNotFoundError:
        return "[未找到 unoconv 命令，请安装 LibreOffice 和 unoconv]"
    except Exception as e:
        return f"[unoconv 转换异常: {str(e)}]"


def read_doc_file_with_textract(path: str) -> str:
    """使用 textract 库读取 .doc 文件（如果可用）"""
    try:
        import textract
        content = textract.process(path).decode('utf-8', errors='ignore')
        return content
    except ImportError:
        return "[textract 不可用，请安装: pip install textract]"
    except Exception as e:
        return f"[textract 转换失败: {str(e)}]"


def read_text_from_any_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()

    # 优先使用扩展名匹配
    if ext == ".csv" or ext == ".tsv":
        with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
            delimiter = "," if ext == ".csv" else "\t"
            reader = csv.reader(f, delimiter=delimiter)
            rows = []
            for row in reader:
                rows.append("\t".join(row))
            return "\n".join(rows)

    if ext == ".doc":
        # 方法1: 尝试使用 win32com (Windows 系统)
        if WIN32_AVAILABLE:
            try:
                content = read_doc_file_with_win32com(path)
                if not content.startswith("[win32com 读取失败"):
                    return content
            except Exception:
                pass
        
        # 方法2: 尝试使用 antiword
        if SUBPROCESS_AVAILABLE:
            try:
                content = read_doc_file_with_antiword(path)
                if not content.startswith("[未找到 antiword 命令") and not content.startswith("[antiword 转换失败"):
                    return content
            except Exception:
                pass
        
        # 方法3: 尝试使用 catdoc
        if SUBPROCESS_AVAILABLE:
            try:
                content = read_doc_file_with_catdoc(path)
                if not content.startswith("[未找到 catdoc 命令") and not content.startswith("[catdoc 转换失败"):
                    return content
            except Exception:
                pass
        
        # 方法4: 尝试使用 unoconv (需要 LibreOffice)
        if SUBPROCESS_AVAILABLE:
            try:
                content = read_doc_file_with_unoconv(path)
                if not content.startswith("[未找到 unoconv 命令") and not content.startswith("[unoconv 转换失败"):
                    return content
            except Exception:
                pass
        
        # 方法5: 尝试使用 textract
        try:
            content = read_doc_file_with_textract(path)
            if not content.startswith("[textract 不可用") and not content.startswith("[textract 转换失败"):
                return content
        except Exception:
            pass
        
        # 方法6: 最后尝试使用 python-docx (可能支持某些格式的 .doc 文件)
        if docx is not None:
            try:
                doc = docx.Document(path)
                out = []
                for p in doc.paragraphs:
                    out.append(p.text)
                # 简单读取表格
                for table in doc.tables:
                    for row in table.rows:
                        out.append(" | ".join([cell.text.strip() for cell in row.cells]))
                content = "\n".join(out)
                if content.strip():  # 如果有内容，认为成功
                    return content
            except Exception:
                pass
        
        # 所有方法都失败，提供安装指导
        install_guide = """
[.doc 文件读取失败]

支持 .doc 文件的几种方案（任选其一）：

1. Windows 系统（推荐）:
   • 安装 Microsoft Word 或 Word Viewer

   • 安装 pywin32: pip install pywin32


2. 安装 antiword（跨平台）:
   • Linux: sudo apt-get install antiword

   • macOS: brew install antiword

   • Windows: 下载 antiword.exe 并添加到 PATH

   • 然后使用: pip install antiword


3. 安装 catdoc（跨平台）:
   • Linux: sudo apt-get install catdoc

   • macOS: brew install catdoc

   • Windows: 下载 catdoc 并添加到 PATH


4. 安装 LibreOffice + unoconv:
   • 安装 LibreOffice

   • 安装 unoconv: pip install unoconv


5. 安装 textract（推荐）:
   • pip install textract


选择一种方案安装后，重新尝试打开 .doc 文件。
        """
        return install_guide.strip()

    if ext == ".docx":
        if docx is None:
            return "[不支持的格式: .docx]\npip install python-docx"
        try:
            doc = docx.Document(path)
            out = []
            for p in doc.paragraphs:
                out.append(p.text)
            # 简单读取表格
            for table in doc.tables:
                for row in table.rows:
                    out.append(" | ".join([cell.text.strip() for cell in row.cells]))
            return "\n".join(out)
        except Exception as e:
            return f"[读取 .docx 失败]\n{str(e)}"

    if ext == ".xlsx":
        if openpyxl is None:
            return "[不支持的格式: .xlsx]\npip install openpyxl"
        try:
            wb = openpyxl.load_workbook(path, data_only=True)
            out_lines = []
            for name in wb.sheetnames:
                ws = wb[name]
                out_lines.append(f"[Sheet: {name}]")
                max_row = ws.max_row or 1
                max_col = ws.max_column or 1
                for r in range(1, max_row + 1):
                    row = []
                    for c in range(1, max_col + 1):
                        v = ws.cell(r, c).value
                        if v is None:
                            row.append("")
                        else:
                            row.append(str(v))
                    out_lines.append("\t".join(row))
                out_lines.append("")
            return "\n".join(out_lines)
        except Exception as e:
            return f"[读取 .xlsx 失败]\n{str(e)}"

    if ext == ".xls":
        if xlrd is None:
            return "[不支持的格式: .xls]\npip install xlrd"
        try:
            book = xlrd.open_workbook(path, formatting_info=True)
            out_lines = []
            for s_index in range(book.nsheets):
                sheet = book.sheet_by_index(s_index)
                out_lines.append(f"[Sheet: {sheet.name}]")
                for r in range(sheet.nrows):
                    row = []
                    for c in range(sheet.ncols):
                        cell = sheet.cell(r, c)
                        val = cell.value
                        # 简单类型转换
                        if cell.ctype == xlrd.XL_CELL_DATE:
                            val = xlrd.xldate_as_tuple(val, book.datemode)
                            val = str(val)
                        elif val is None:
                            val = ""
                        else:
                            val = str(val)
                        row.append(val)
                    out_lines.append("\t".join(row))
                out_lines.append("")
            return "\n".join(out_lines)
        except Exception as e:
            return f"[读取 .xls 失败]\n{str(e)}"

    # 默认按文本读取
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"[读取文件失败: {path}]\n{str(e)}"


def search_text_files_recursive(root_dir: str) -> List[str]:
    """优化版本：添加文件过滤和限制"""
    result = []
    total_size = 0
    file_count = 0
    max_files = 1000  # 最大文件数量限制
    
    for dirpath, dirnames, filenames in os.walk(root_dir):
        # 跳过一些常见的大文件目录
        skip_dirs = {'.git', '__pycache__', 'node_modules', '.idea', 'target', 'build'}
        dirnames[:] = [d for d in dirnames if d not in skip_dirs]
        
        for fn in filenames:
            if file_count >= max_files:
                result.append(f"[警告: 已达到最大文件数量限制 {max_files}，停止扫描]")
                return result
                
            ext = os.path.splitext(fn)[1].lower()
            if ext in ALLOWED_TEXT_EXTS:
                full_path = os.path.join(dirpath, fn)
                try:
                    file_size = os.path.getsize(full_path)
                    # 跳过超大文件
                    if file_size > FILE_SIZE_LIMIT:
                        continue
                    total_size += file_size
                    # 限制总大小
                    if total_size > 500 * 1024 * 1024:  # 500MB
                        result.append(f"[警告: 总文件大小超过500MB，停止扫描]")
                        return result
                    result.append(full_path)
                    file_count += 1
                except OSError:
                    continue
    
    return result


def guess_display_name_from_content(content: str) -> str:
    # 优先使用第一行非空作为显示名
    for line in content.splitlines():
        if line.strip():
            return line.strip()[:80]
    # 否则取前80个字符
    return content.strip()[:80]


def int_to_excel_col(n: int) -> str:
    s = ""
    while n > 0:
        n -= 1
        s = chr(65 + (n % 26)) + s
        n //= 26
    return s


@dataclass
class DataItem:
    """
    kind: manual | text_file | siyuan_id | siyuan_multi_id | folder
    payload:
      ◦ manual: {text}

      ◦ text_file: {file_path, content, size, display_name}

      ◦ siyuan_id: {id, title, content}

      ◦ siyuan_multi_id: {id, title, content}

      ◦ folder: {folder_path, files: [{file_path, content, size, display_name}]}

    """
    kind: str
    payload: Dict[str, Any]
    _id: str = field(default_factory=lambda: uuid4().hex)

    def display_label(self) -> str:
        if self.kind == "manual":
            preview = self.payload.get("text", "").strip().splitlines()[0][:60]
            return f"手动提示词: {preview if preview else '(空)'}"
        elif self.kind == "text_file":
            size = self.payload.get("size", 0)
            # 优先显示文件名（含扩展名）
            name = os.path.basename(self.payload.get("file_path", "文件"))
            size_str = f"{size:,} bytes" if size > 1000 else f"{size} bytes"
            return f"文本文件: {name} ({size_str})"
        elif self.kind == "siyuan_id":
            # 优先显示路径标题，后接ID
            title = self.payload.get("title") or "无标题"
            tid = self.payload.get("id", "")
            return f"SiYuan文档: {title} [{tid}]"
        elif self.kind == "siyuan_multi_id":
            title = self.payload.get("title") or "无标题"
            tid = self.payload.get("id", "")
            return f"SiYuan文档(多ID): {title} [{tid}]"
        elif self.kind == "folder":
            path = self.payload.get("folder_path", "")
            count = len(self.payload.get("files", []))
            return f"文件夹: {path} ({count} 个文件)"
        else:
            return f"未知类型: {self.kind}"

    def to_dict(self) -> Dict[str, Any]:
        return {"kind": self.kind, "payload": self.payload, "_id": self._id}

    @staticmethod
    def from_dict(d: Dict[str, Any]) -> "DataItem":
        return DataItem(kind=d["kind"], payload=d["payload"], _id=d["_id"])


class FileProcessor(QtCore.QObject):
    """专门处理文件读取的类"""
    file_processed = QtCore.pyqtSignal(object)  # DataItem
    progress_updated = QtCore.pyqtSignal(int, int, str)
    finished = QtCore.pyqtSignal()
    
    def __init__(self, file_paths: List[str]):
        super().__init__()
        self.file_paths = file_paths
        self._cancelled = False
        
    def cancel(self):
        self._cancelled = True
        
    def process_files(self):
        """在后台线程中处理文件"""
        total = len(self.file_paths)
        
        for idx, file_path in enumerate(self.file_paths, 1):
            if self._cancelled:
                break
                
            try:
                size = os.path.getsize(file_path)
                if size > FILE_SIZE_LIMIT:
                    content = f"[文件过大，未读取: {os.path.basename(file_path)} ({size} bytes)]"
                else:
                    content = read_text_from_any_file(file_path)
                
                display_name = None
                if not content.startswith("[文件过大，未读取") and not content.startswith("[读取失败"):
                    display_name = guess_display_name_from_content(content)
                
                payload = {
                    "file_path": file_path,
                    "content": content,
                    "size": size,
                    "display_name": display_name or os.path.basename(file_path)
                }
                item = DataItem(kind="text_file", payload=payload)
                
                self.file_processed.emit(item)
                self.progress_updated.emit(idx, total, f"处理文件 {idx}/{total}")
                
            except Exception as e:
                # 创建错误项
                payload = {
                    "file_path": file_path,
                    "content": f"[处理文件时出错: {str(e)}]",
                    "size": 0,
                    "display_name": os.path.basename(file_path)
                }
                item = DataItem(kind="text_file", payload=payload)
                self.file_processed.emit(item)
        
        self.finished.emit()


class SiYuanHelper(QtCore.QObject):
    status_signal = QtCore.pyqtSignal(str)
    progress_signal = QtCore.pyqtSignal(int, int, str)

    def __init__(self, cfg: Dict[str, Any]):
        super().__init__()
        self.api_base_url = cfg.get("api_base_url", "").strip().rstrip("/")
        self.api_token = cfg.get("api_token", "")
        self.timeout = int(cfg.get("timeout", "8"))

    def get_document(self, doc_id: str) -> Tuple[Optional[str], Optional[str]]:
        """
        使用 /api/export/exportMdContent 接口拉取文档内容
        返回: (title, content)
        """
        if not self.api_base_url or not self.api_token:
            return None, None

        headers = {
            "Authorization": f"Token {self.api_token}",
            "Content-Type": "application/json"
        }

        payload = {"id": doc_id}

        try:
            resp = requests.post(
                f"{self.api_base_url}/api/export/exportMdContent",
                headers=headers,
                json=payload,
                timeout=self.timeout
            )
            if resp.status_code != 200:
                return None, None
            j = resp.json()
            if j.get("code") != 0:
                return None, None
            data = j.get("data", {})
            if not isinstance(data, dict):
                return None, None

            # 成功时，data 结构示例：
            # {
            #   "hPath": "...",
            #   "content": "文档 markdown 内容"
            # }
            title = data.get("hPath") or doc_id
            content = data.get("content", "")
            return title, content
        except Exception:
            return None, None


class SavedPromptsManager:
    def __init__(self, cfg_path: str = DEFAULT_CONFIG_PATH):
        self.cfg_path = cfg_path
        self._lock = threading.Lock()
        self._data = {"manual_prompts": []}  # 每个项: {text, note, created, updated}
        self._ensure_file()
        self._load()

    def _ensure_file(self):
        if not os.path.exists(self.cfg_path):
            try:
                with open(self.cfg_path, "w", encoding="utf-8") as f:
                    json.dump(self._data, f, ensure_ascii=False, indent=2)
            except Exception:
                pass

    def _load(self):
        with self._lock:
            if os.path.exists(self.cfg_path):
                try:
                    with open(self.cfg_path, "r", encoding="utf-8") as f:
                        self._data = json.load(f)
                    if "manual_prompts" not in self._data:
                        self._data["manual_prompts"] = []
                except Exception:
                    self._data = {"manual_prompts": []}

    def _save(self):
        with self._lock:
            try:
                with open(self.cfg_path, "w", encoding="utf-8") as f:
                    json.dump(self._data, f, ensure_ascii=False, indent=2)
            except Exception:
                pass

    def list_prompts(self) -> List[Dict[str, Any]]:
        with self._lock:
            return list(self._data.get("manual_prompts", []))

    def add_or_update_prompt(self, text: str, note: str = ""):
        text = text.strip()
        if not text:
            return
        with self._lock:
            existing = next((p for p in self._data["manual_prompts"] if p["text"] == text), None)
            if existing:
                existing["updated"] = now_label()
                existing["note"] = note.strip() if note else ""
            else:
                self._data["manual_prompts"].append({
                    "text": text,
                    "note": note.strip() if note else "",
                    "created": now_label(),
                    "updated": now_label()
                })
        self._save()

    def delete_prompt(self, text: str):
        with self._lock:
            self._data["manual_prompts"] = [p for p in self._data["manual_prompts"] if p["text"] != text]
        self._save()


class AddArea(QtWidgets.QWidget):
    filesDropped = QtCore.pyqtSignal(list)  # list[str]
    folderSelected = QtCore.pyqtSignal(str)  # str
    noteIdsSubmitted = QtCore.pyqtSignal(list)  # list[str]
    manualSubmitted = QtCore.pyqtSignal(str)  # str

    def __init__(self, saved_manager: Optional[SavedPromptsManager] = None, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.saved_manager = saved_manager or SavedPromptsManager()
        self.init_ui()

    def init_ui(self):
        v = QtWidgets.QVBoxLayout(self)
        v.setContentsMargins(6, 6, 6, 6)
        v.setSpacing(8)

        title = QtWidgets.QLabel("左侧数据源（拖入或输入）")
        v.addWidget(title)

        # 手动输入
        manual_box = QtWidgets.QGroupBox("1) 手动提示词（可保存）")
        manual_layout = QtWidgets.QVBoxLayout(manual_box)

        saved_layout = QtWidgets.QHBoxLayout()
        self.saved_combo = QtWidgets.QComboBox()
        self.saved_combo.setMinimumWidth(220)
        self.refresh_saved_combo()
        self.saved_combo.currentTextChanged.connect(self.on_saved_combo_changed)
        self.load_saved_btn = QtWidgets.QPushButton("加载到编辑器")
        self.save_btn = QtWidgets.QPushButton("保存/更新此条")
        self.delete_btn = QtWidgets.QPushButton("删除此条")
        saved_layout.addWidget(QtWidgets.QLabel("已保存提示词:"))
        saved_layout.addWidget(self.saved_combo, 1)
        saved_layout.addWidget(self.load_saved_btn)
        saved_layout.addWidget(self.save_btn)
        saved_layout.addWidget(self.delete_btn)
        manual_layout.addLayout(saved_layout)

        self.manual_edit = QtWidgets.QPlainTextEdit()
        self.manual_edit.setPlaceholderText("在此输入提示词文本，支持多行。可保存到右侧列表。")
        self.manual_edit.setMinimumHeight(120)
        manual_layout.addWidget(self.manual_edit)
        btns = QtWidgets.QHBoxLayout()
        add_manual_btn = QtWidgets.QPushButton("添加到右侧")
        btns.addWidget(add_manual_btn)
        manual_layout.addLayout(btns)
        add_manual_btn.clicked.connect(self._on_add_manual)
        self.load_saved_btn.clicked.connect(self._on_load_saved)
        self.save_btn.clicked.connect(self._on_save_prompt)
        self.delete_btn.clicked.connect(self._on_delete_saved)
        v.addWidget(manual_box)

        # 思源笔记 ID 输入
        id_box = QtWidgets.QGroupBox("2) 思源笔记文档ID（可多行）")
        id_layout = QtWidgets.QVBoxLayout(id_box)
        self.id_edit = QtWidgets.QPlainTextEdit()
        self.id_edit.setPlaceholderText("输入多个思源笔记文档ID，例如：\n20251108093753-hph79vi\n20251020103018-rz9r4ge")
        self.id_edit.setMinimumHeight(100)
        id_layout.addWidget(self.id_edit)
        add_id_btn = QtWidgets.QPushButton("解析ID并添加")
        add_id_btn.clicked.connect(self._on_add_ids)
        id_layout.addWidget(add_id_btn)
        v.addWidget(id_box)

        # 拖入文本文件
        drop_box = QtWidgets.QGroupBox("3) 拖入文本文件到此区域（支持 txt/md/json/yml/xml/html/js/py 以及 office: doc/docx/xls/xlsx 等）")
        drop_layout = QtWidgets.QVBoxLayout(drop_box)
        self.drop_label = QtWidgets.QLabel("将文本或 office 文档拖入此处（.txt/.md/.doc/.docx/.xls/.xlsx/.csv/.json/.yml/.xml/.html/.js/.py 等）")
        self.drop_label.setWordWrap(True)
        self.drop_label.setStyleSheet("border: 1px dashed #aaa; padding: 20px;")
        self.drop_label.setAlignment(QtCore.Qt.AlignCenter)
        drop_layout.addWidget(self.drop_label)
        v.addWidget(drop_box)

        # 文件夹选择
        folder_box = QtWidgets.QGroupBox("4) 选择文件夹（无限递归，扫描上述支持格式）")
        folder_layout = QtWidgets.QHBoxLayout(folder_box)
        self.folder_path = QtWidgets.QLineEdit()
        self.folder_path.setPlaceholderText("选择目录，将自动扫描子目录中的文本/office 文件")
        choose_btn = QtWidgets.QPushButton("选择...")
        scan_btn = QtWidgets.QPushButton("扫描并添加")
        folder_layout.addWidget(self.folder_path)
        folder_layout.addWidget(choose_btn)
        folder_layout.addWidget(scan_btn)
        choose_btn.clicked.connect(self._choose_folder)
        scan_btn.clicked.connect(self._scan_folder)
        v.addWidget(folder_box)

        v.addStretch(1)

    def refresh_saved_combo(self):
        self.saved_combo.blockSignals(True)
        self.saved_combo.clear()
        prompts = self.saved_manager.list_prompts()
        if not prompts:
            self.saved_combo.addItem("(无已保存的提示词)")
        else:
            for p in prompts:
                text = (p.get("text") or "")
                note = p.get("note", "").strip()
                
                # 优先显示备注，如果没有备注则显示文本预览
                if note:
                    display = note
                else:
                    preview = text.strip().splitlines()[0][:60] if text.strip() else "(空)"
                    display = preview
                
                self.saved_combo.addItem(display, userData={
                    "text": text, 
                    "note": note,
                    "created": p.get("created", ""),
                    "updated": p.get("updated", "")
                })
        self.saved_combo.blockSignals(False)

    def on_saved_combo_changed(self, text: str):
        pass

    def dragEnterEvent(self, event: QtGui.QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            super().dragEnterEvent(event)

    def dragMoveEvent(self, event: QtGui.QDragMoveEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            super().dragMoveEvent(event)

    def dropEvent(self, event: QtGui.QDropEvent):
        files = []
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if os.path.isdir(path):
                files.append(path)
            elif os.path.isfile(path):
                files.append(path)
        if files:
            folders = [f for f in files if os.path.isdir(f)]
            plain_files = [f for f in files if os.path.isfile(f)]
            if folders:
                self.folderSelected.emit(folders[0])
            if plain_files:
                self.filesDropped.emit(plain_files)
            event.acceptProposedAction()
        else:
            super().dropEvent(event)

    def _on_add_manual(self):
        text = self.manual_edit.toPlainText().strip()
        if not text:
            QtWidgets.QMessageBox.information(self, "提示", "请先输入提示词内容")
            return
        self.manualSubmitted.emit(text)

    def _on_add_ids(self):
        raw = self.id_edit.toPlainText()
        ids = normalize_id_list(raw)
        if not ids:
            QtWidgets.QMessageBox.information(self, "提示", "未识别到有效的文档ID")
            return
        self.noteIdsSubmitted.emit(ids)

    def _choose_folder(self):
        path = QtWidgets.QFileDialog.getExistingDirectory(self, "选择文件夹")
        if path:
            self.folder_path.setText(path)

    def _scan_folder(self):
        path = self.folder_path.text().strip()
        if not path or not os.path.isdir(path):
            QtWidgets.QMessageBox.information(self, "提示", "请先选择有效目录")
            return
        self.folderSelected.emit(path)

    def _on_load_saved(self):
        idx = self.saved_combo.currentIndex()
        if idx < 0:
            return
        data = self.saved_combo.itemData(idx)
        if not isinstance(data, dict):
            return
        text = data.get("text", "")
        if text:
            self.manual_edit.setPlainText(text)

    def _on_save_prompt(self):
        text = self.manual_edit.toPlainText().strip()
        if not text:
            QtWidgets.QMessageBox.information(self, "提示", "请在编辑器中输入或加载要保存的提示词")
            return
        
        # 检查是否已存在相同的提示词
        prompts = self.saved_manager.list_prompts()
        existing_prompt = next((p for p in prompts if p.get("text") == text), None)
        existing_note = existing_prompt.get("note", "").strip() if existing_prompt else ""
        
        # 询问是否添加/修改备注
        if existing_note:
            msg = f"是否修改备注？\n当前备注：{existing_note}"
        else:
            msg = "是否为此提示词添加备注？"
        
        ret = QtWidgets.QMessageBox.question(
            self, "保存提示词", msg,
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        
        note = existing_note  # 默认保留原有备注
        if ret == QtWidgets.QMessageBox.Yes:
            # 弹出输入对话框获取备注
            note_text, ok = QtWidgets.QInputDialog.getText(
                self, 
                "添加/修改备注", 
                "请输入备注:",
                QtWidgets.QLineEdit.Normal,
                existing_note  # 显示原有备注（如果有）
            )
            if ok:
                # 用户确认输入，使用新输入的备注（可能为空字符串，表示清空备注）
                note = note_text.strip()
            # 如果取消，note保持为原有备注，不修改
        
        self.saved_manager.add_or_update_prompt(text, note)
        self.refresh_saved_combo()
        QtWidgets.QMessageBox.information(self, "成功", "提示词已保存/更新")

    def _on_delete_saved(self):
        idx = self.saved_combo.currentIndex()
        if idx < 0:
            return
        data = self.saved_combo.itemData(idx)
        if not isinstance(data, dict):
            return
        text = data.get("text", "")
        if not text:
            return
        ret = QtWidgets.QMessageBox.question(self, "确认", "确定要删除这条已保存的提示词吗？")
        if ret == QtWidgets.QMessageBox.Yes:
            self.saved_manager.delete_prompt(text)
            self.refresh_saved_combo()


class ReorderListWidget(QtWidgets.QListWidget):
    itemMoved = QtCore.pyqtSignal()
    itemDeleted = QtCore.pyqtSignal()  # 删除项时发出
    itemDuplicated = QtCore.pyqtSignal()  # 复制项时发出
    itemContentRequested = QtCore.pyqtSignal(object)  # 请求复制内容时发出，传递 DataItem（单选）
    itemsContentRequested = QtCore.pyqtSignal(list)  # 请求复制内容时发出，传递 DataItem 列表（多选）

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setIconSize(QtCore.QSize(18, 18))
        self.setSelectionMode(QtWidgets.QAbstractItemView.ExtendedSelection)
        self.setDragEnabled(True)
        self.setDragDropMode(QtWidgets.QAbstractItemView.InternalMove)
        self.setDefaultDropAction(QtCore.Qt.MoveAction)
        self.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        self.customContextMenuRequested.connect(self._on_context_menu)

    def mimeTypes(self):
        return ["text/plain", "application/x-qabstractitemmodeldatalist"]

    def dropEvent(self, e: QtGui.QDropEvent):
        super().dropEvent(e)
        self.itemMoved.emit()

    def _on_context_menu(self, pos):
        item = self.itemAt(pos)
        if not item:
            return
        menu = QtWidgets.QMenu(self)
        act_del = menu.addAction("删除")
        act_dup = menu.addAction("复制")
        act_up = menu.addAction("上移")
        act_down = menu.addAction("下移")
        act_top = menu.addAction("置顶")
        act_bottom = menu.addAction("置底")
        act = menu.exec_(self.mapToGlobal(pos))
        if act is None:
            return
        
        # 获取所有选中的项
        selected_items = self.selectedItems()
        if not selected_items:
            selected_items = [item]  # 如果没有选中项，使用右键点击的项
        
        if act == act_del:
            # 删除前确认
            count = len(selected_items)
            msg = f"确定要删除这 {count} 个项目吗？此操作不可恢复。" if count > 1 else "确定要删除这个项目吗？此操作不可恢复。"
            ret = QtWidgets.QMessageBox.question(
                self.parent(), "确认", msg,
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
            )
            if ret != QtWidgets.QMessageBox.Yes:
                return
            # 删除所有选中的项
            for it in selected_items:
                row = self.row(it)
                if row >= 0:
                    self.takeItem(row)
            self.itemDeleted.emit()  # 发出删除信号
        elif act == act_dup:
            # 复制所有选中项的内容到剪贴板
            items_to_copy = []
            for it in selected_items:
                data = it.data(QtCore.Qt.UserRole)
                if data:
                    try:
                        item = DataItem.from_dict(data)
                        items_to_copy.append(item)
                    except Exception:
                        pass
            if items_to_copy:
                # 发出多选复制信号
                self.itemsContentRequested.emit(items_to_copy)
        elif act == act_up:
            # 上移操作只对单个项有效
            row = self.row(item)
            if row > 0:
                self.insertItem(row - 1, self.takeItem(row))
                self.itemMoved.emit()
        elif act == act_down:
            # 下移操作只对单个项有效
            row = self.row(item)
            if row < self.count() - 1:
                self.insertItem(row + 1, self.takeItem(row))
                self.itemMoved.emit()
        elif act == act_top:
            # 置顶操作只对单个项有效
            row = self.row(item)
            it = self.takeItem(row)
            if it:
                self.insertItem(0, it)
                self.itemMoved.emit()
        elif act == act_bottom:
            # 置底操作只对单个项有效
            row = self.row(item)
            it = self.takeItem(row)
            if it:
                self.addItem(it)
                self.itemMoved.emit()


class SettingsDialog(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("设置 - 思源笔记 API")
        self.resize(520, 260)
        self.cfg: Dict[str, Any] = {}
        self.init_ui()
        self.load_from_default()

    def init_ui(self):
        v = QtWidgets.QVBoxLayout(self)

        api_box = QtWidgets.QGroupBox("思源笔记 HTTP API 配置")
        g = QtWidgets.QFormLayout(api_box)
        self.api_url = QtWidgets.QLineEdit()
        self.api_url.setPlaceholderText("http://127.0.0.1:6806")
        self.api_token = QtWidgets.QLineEdit()
        self.api_token.setEchoMode(QtWidgets.QLineEdit.Password)
        self.api_token.setPlaceholderText("请输入 Token，例如：oc8**lp")
        g.addRow("API Base URL:", self.api_url)
        g.addRow("API Token:", self.api_token)
        v.addWidget(api_box)

        note = QtWidgets.QLabel(
            "说明：\n"
            "- 需要在思源笔记中开启服务（默认端口 6806），并获取 API Token。\n"
            "- 在思源笔记：'关于-高级设置-Token'复制 Token。\n"
            "- 本工具使用 /api/export/exportMdContent 接口拉取文档 Markdown 内容。\n"
            "- 确保思源笔记服务启动成功且可达。\n"
        )
        note.setWordWrap(True)
        v.addWidget(note)

        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Ok | QtWidgets.QDialogButtonBox.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        v.addWidget(btns)

    def get_config(self) -> Dict[str, Any]:
        return {
            "api_base_url": self.api_url.text().strip(),
            "api_token": self.api_token.text().strip(),
            "timeout": "8"
        }

    def set_config(self, cfg: Dict[str, Any]):
        self.api_url.setText(cfg.get("api_base_url", ""))
        self.api_token.setText(cfg.get("api_token", ""))

    def load_from_default(self):
        if os.path.exists(DEFAULT_CONFIG_PATH):
            try:
                with open(DEFAULT_CONFIG_PATH, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
                self.set_config(cfg)
            except Exception:
                pass

    def accept(self):
        self.cfg = self.get_config()
        cfg_data = {}
        if os.path.exists(DEFAULT_CONFIG_PATH):
            try:
                with open(DEFAULT_CONFIG_PATH, "r", encoding="utf-8") as f:
                    cfg_data = json.load(f)
            except Exception:
                cfg_data = {}
        # 只更新 API 相关字段
        cfg_data["api_base_url"] = self.cfg["api_base_url"]
        cfg_data["api_token"] = self.cfg["api_token"]
        cfg_data["timeout"] = "8"
        try:
            with open(DEFAULT_CONFIG_PATH, "w", encoding="utf-8") as f:
                json.dump(cfg_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "警告", f"保存配置失败：\n{str(e)}")
        super().accept()


class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(APP_NAME)
        self.resize(1100, 680)
        self.items: List[DataItem] = []
        self.siyuan_helper: Optional[SiYuanHelper] = None
        self.saved_manager = SavedPromptsManager()
        self.file_processor: Optional[FileProcessor] = None
        self.processing_files = False

        self.init_ui()
        self.load_config()
        self.load_list_data()  # 加载保存的列表数据

    def init_ui(self):
        central = QtWidgets.QWidget(self)
        self.setCentralWidget(central)
        h = QtWidgets.QHBoxLayout(central)
        h.setContentsMargins(8, 8, 8, 8)
        h.setSpacing(8)

        # 左侧：AddArea
        self.add_area = AddArea(saved_manager=self.saved_manager)
        self.add_area.filesDropped.connect(self.on_files_dropped)
        self.add_area.folderSelected.connect(self.on_folder_selected)
        self.add_area.noteIdsSubmitted.connect(self.on_note_ids_submitted)
        self.add_area.manualSubmitted.connect(self.on_manual_submitted)
        w_left = QtWidgets.QWidget()
        v_left = QtWidgets.QVBoxLayout(w_left)
        v_left.addWidget(self.add_area, 1)

        # 右侧：List + Controls
        right = QtWidgets.QWidget()
        v_right = QtWidgets.QVBoxLayout(right)
        v_right.setContentsMargins(0, 0, 0, 0)
        list_label = QtWidgets.QLabel("数据项列表（可拖动排序）")
        v_right.addWidget(list_label)
        self.list = ReorderListWidget()
        self.list.itemMoved.connect(self.on_list_item_moved)  # 连接移动信号
        self.list.itemDeleted.connect(self.on_list_item_moved)  # 删除后保存
        self.list.itemContentRequested.connect(self.on_copy_item_content)  # 复制单个内容
        self.list.itemsContentRequested.connect(self.on_copy_items_content)  # 复制多个内容
        v_right.addWidget(self.list, 1)

        controls = QtWidgets.QHBoxLayout()
        self.merge_btn = QtWidgets.QPushButton("合并并复制")
        self.clear_btn = QtWidgets.QPushButton("清空列表")
        self.cancel_btn = QtWidgets.QPushButton("取消处理")
        self.cancel_btn.setVisible(False)
        self.settings_btn = QtWidgets.QPushButton("设置")
        self.progress = QtWidgets.QProgressBar()
        self.progress.setRange(0, 100)
        self.progress.setValue(0)
        controls.addWidget(self.merge_btn)
        controls.addWidget(self.clear_btn)
        controls.addWidget(self.cancel_btn)
        controls.addWidget(self.settings_btn)
        controls.addWidget(self.progress)
        v_right.addLayout(controls)
        
        self.log_view = QtWidgets.QPlainTextEdit()
        self.log_view.setReadOnly(True)
        self.log_view.setMaximumHeight(100)
        v_right.addWidget(self.log_view)

        splitter = QtWidgets.QSplitter()
        splitter.addWidget(w_left)
        splitter.addWidget(right)
        splitter.setSizes([420, 680])
        h.addWidget(splitter, 1)

        # 状态栏
        self.statusBar().showMessage("就绪")

        # 事件
        self.merge_btn.clicked.connect(self.on_merge_and_copy)
        self.clear_btn.clicked.connect(self.on_clear_list)
        self.cancel_btn.clicked.connect(self.on_cancel_processing)
        self.settings_btn.clicked.connect(self.on_settings)

    def load_config(self):
        cfg = {"api_base_url": "", "api_token": "", "timeout": "8"}
        if os.path.exists(DEFAULT_CONFIG_PATH):
            try:
                with open(DEFAULT_CONFIG_PATH, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
            except Exception:
                pass
        self.siyuan_helper = SiYuanHelper(cfg)

    def log(self, msg: str):
        self.log_view.appendPlainText(f"[{now_label()}] {msg}")

    def set_processing_state(self, processing: bool):
        """设置处理状态，更新UI"""
        self.processing_files = processing
        self.cancel_btn.setVisible(processing)
        self.merge_btn.setEnabled(not processing)
        self.clear_btn.setEnabled(not processing)
        self.settings_btn.setEnabled(not processing)
        self.add_area.setEnabled(not processing)

    # 左侧回调
    def on_manual_submitted(self, text: str):
        item = DataItem(kind="manual", payload={"text": text})
        self.add_item(item)

    def on_files_dropped(self, file_paths: List[str]):
        if self.processing_files:
            QtWidgets.QMessageBox.information(self, "提示", "当前正在处理文件，请等待完成或取消")
            return
            
        if not file_paths:
            return
            
        # 检查文件数量限制
        if len(file_paths) > 100:
            ret = QtWidgets.QMessageBox.question(
                self, "确认", 
                f"检测到 {len(file_paths)} 个文件，数量较多可能影响性能。\n是否继续处理？",
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
            )
            if ret != QtWidgets.QMessageBox.Yes:
                return
        
        self.set_processing_state(True)
        self.statusBar().showMessage(f"正在处理 {len(file_paths)} 个文件...")
        
        # 使用后台线程处理文件
        self.file_processor = FileProcessor(file_paths)
        self.file_processor.file_processed.connect(self.add_item)
        self.file_processor.progress_updated.connect(self._set_progress)
        self.file_processor.finished.connect(self.on_file_processing_finished)
        
        # 启动处理线程
        thread = threading.Thread(target=self.file_processor.process_files, daemon=True)
        thread.start()

    def on_folder_selected(self, folder: str):
        if self.processing_files:
            QtWidgets.QMessageBox.information(self, "提示", "当前正在处理文件，请等待完成或取消")
            return
            
        self.statusBar().showMessage("正在扫描文件夹...")
        QtWidgets.QApplication.processEvents()
        
        files = search_text_files_recursive(folder)
        if not files:
            QtWidgets.QMessageBox.information(self, "提示", f"未在目录 {folder} 及其子目录中找到支持的文件")
            return
            
        # 显示扫描结果确认
        file_count = len([f for f in files if not f.startswith("[警告:")])
        if file_count > 0:
            ret = QtWidgets.QMessageBox.question(
                self, "确认", 
                f"扫描到 {file_count} 个文件。是否开始处理？",
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
            )
            if ret == QtWidgets.QMessageBox.Yes:
                self.on_files_dropped([f for f in files if not f.startswith("[警告:")])
        else:
            QtWidgets.QMessageBox.information(self, "提示", "未找到符合条件的文件")

    def on_note_ids_submitted(self, ids: List[str]):
        if not ids:
            return
        if not self._check_siyuan_available():
            return
        self._fetch_siyuan_ids_blocking(ids)

    def on_file_processing_finished(self):
        """文件处理完成回调"""
        self.set_processing_state(False)
        self._set_progress(0, 1, "文件处理完成")
        self.log(f"文件处理完成，共处理 {self.list.count()} 个项目")

    def on_cancel_processing(self):
        """取消文件处理"""
        if self.file_processor:
            self.file_processor.cancel()
            self.log("已取消文件处理")
        self.set_processing_state(False)
        self._set_progress(0, 1, "已取消")

    def _check_siyuan_available(self) -> bool:
        # 检查 API 是否配置
        if self.siyuan_helper.api_base_url and self.siyuan_helper.api_token:
            self.log("将使用 SiYuan HTTP API 获取数据")
            return True
        QtWidgets.QMessageBox.information(self, "提示",
                                          "未配置或不可用的思源笔记连接。\n"
                                          "请点击'设置'，配置 API URL 和 Token。")
        return False

    def _fetch_siyuan_ids_blocking(self, ids: List[str]):
        """
        在后台线程按顺序获取并添加到列表
        """
        def worker():
            total = len(ids)
            for idx, doc_id in enumerate(ids, 1):
                title, content = self.siyuan_helper.get_document(doc_id)
                if content and content.strip():
                    payload = {"id": doc_id, "title": title or "无标题", "content": content}
                    item = DataItem(kind="siyuan_id", payload=payload)
                else:
                    payload = {"id": doc_id, "title": "未能获取", "content": "[未能获取内容]"}
                    item = DataItem(kind="siyuan_id", payload=payload)

                self._append_item_in_main(item)
                self._update_progress(idx, total, f"已添加 {idx}/{total}")
            self._update_progress(total, total, "完成")
        threading.Thread(target=worker, daemon=True).start()

    def _append_item_in_main(self, item: DataItem):
        QtCore.QMetaObject.invokeMethod(self, "_append_item", QtCore.Qt.QueuedConnection, QtCore.Q_ARG(object, item))

    @QtCore.pyqtSlot(object)
    def _append_item(self, item: DataItem):
        self.add_item(item)

    def _update_progress(self, cur: int, total: int, msg: str):
        QtCore.QMetaObject.invokeMethod(self, "_set_progress", QtCore.Qt.QueuedConnection,
                                         QtCore.Q_ARG(int, cur), QtCore.Q_ARG(int, total), QtCore.Q_ARG(str, msg))

    @QtCore.pyqtSlot(int, int, str)
    def _set_progress(self, cur: int, total: int, msg: str):
        if total <= 0:
            self.progress.setRange(0, 1)
            self.progress.setValue(0)
        else:
            self.progress.setRange(0, total)
            self.progress.setValue(cur)
        self.statusBar().showMessage(msg)

    def add_item(self, item: DataItem):
        self.items.append(item)
        # 列表显示
        icon = self._icon_for_kind(item.kind)
        text = item.display_label()
        lw_item = QtWidgets.QListWidgetItem(icon, text)
        lw_item.setData(QtCore.Qt.UserRole, item.to_dict())
        self.list.addItem(lw_item)
        self.log(f"添加: {text}")
        self.save_list_data()  # 保存列表数据

    def _icon_for_kind(self, kind: str) -> QtGui.QIcon:
        pm = QtGui.QPixmap(16, 16)
        if kind == "manual":
            pm.fill(QtGui.QColor("#3f8efc"))
        elif kind == "text_file":
            pm.fill(QtGui.QColor("#4db6ac"))
        elif kind.startswith("siyuan"):
            pm.fill(QtGui.QColor("#9575cd"))
        elif kind == "folder":
            pm.fill(QtGui.QColor("#f06292"))
        else:
            pm.fill(QtGui.QColor("#9e9e9e"))
        return QtGui.QIcon(pm)

    def collect_items_from_list(self) -> List[DataItem]:
        items = []
        for i in range(self.list.count()):
            d = self.list.item(i).data(QtCore.Qt.UserRole)
            if d:
                items.append(DataItem.from_dict(d))
        return items

    def on_merge_and_copy(self):
        items = self.collect_items_from_list()
        if not items:
            QtWidgets.QMessageBox.information(self, "提示", "请先添加数据项")
            return

        merged_lines: List[str] = []
        sep = "\n" + "-" * 80 + "\n"

        try:
            for idx, item in enumerate(items, 1):
                # 跳过空内容项
                if not item.payload:
                    merged_lines.append(f"【无效数据项】\n")
                    continue

                if item.kind == "manual":
                    text = item.payload.get("text", "").strip()
                    if text:
                        merged_lines.append(f"【手动提示词】\n{text}\n")
                    else:
                        merged_lines.append(f"【手动提示词（空）】\n")

                elif item.kind == "text_file":
                    # 修复核心错误：self.payload → item.payload
                    file_path = item.payload.get("file_path", "未知文件")
                    file_name = os.path.basename(file_path)
                    content = item.payload.get("content", "").strip()
                    
                    if content:
                        merged_lines.append(f"【{file_name}】\n{content}\n")
                    else:
                        merged_lines.append(f"【{file_name}（空内容）】\n")

                elif item.kind.startswith("siyuan"):
                    title = item.payload.get("title", "无标题").strip()
                    doc_id = item.payload.get("id", "").strip()
                    content = item.payload.get("content", "").strip()
                    
                    if content:
                        merged_lines.append(f"【{title} [{doc_id}]】\n{content}\n")
                    else:
                        merged_lines.append(f"【{title} [{doc_id}]（空内容）】\n")

                elif item.kind == "folder":
                    folder_path = item.payload.get("folder_path", "未知文件夹")
                    file_count = len(item.payload.get("files", []))
                    merged_lines.append(f"【文件夹：{folder_path}（{file_count}个文件）】\n")
                    # 拼接文件夹内所有文件内容
                    for file_info in item.payload.get("files", []):
                        file_name = os.path.basename(file_info.get("file_path", "未知文件"))
                        file_content = file_info.get("content", "").strip()
                        if file_content:
                            merged_lines.append(f"  ├─ {file_name}\n{file_content}\n")
                        else:
                            merged_lines.append(f"  ├─ {file_name}（空内容）\n")

                else:
                    merged_lines.append(f"【未知类型：{item.kind}】\n")

                # 添加分隔符（最后一项不添加）
                if idx < len(items):
                    merged_lines.append(sep)

            # 合并所有内容（处理超长文本）
            merged_text = "".join(merged_lines)
            
            # 复制到剪贴板（处理空文本）
            if merged_text.strip():
                clipboard = QtWidgets.QApplication.clipboard()
                clipboard.setText(merged_text)
                self.log("已合并并复制到剪贴板")
            else:
                self.log("合并后无有效内容")
                QtWidgets.QMessageBox.warning(self, "提示", "所有数据项均为空内容，无法复制")

        except Exception as e:
            # 捕获所有异常，避免闪退，同时输出错误信息便于排查
            error_msg = f"合并复制失败：{str(e)}\n{traceback.format_exc()[:500]}"
            self.log(error_msg)
            QtWidgets.QMessageBox.critical(self, "错误", f"合并复制时发生异常，已记录日志：\n{str(e)}")

    def on_clear_list(self):
        if self.list.count() == 0:
            return
        ret = QtWidgets.QMessageBox.question(
            self, "确认", "确定要清空列表吗？此操作不可恢复。",
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        if ret != QtWidgets.QMessageBox.Yes:
            return
        self.items.clear()
        self.list.clear()
        self.log("清空列表")
        self.save_list_data()  # 保存清空后的状态
    
    def on_list_item_moved(self):
        """列表项移动/删除/复制后保存，并同步 self.items"""
        # 同步 self.items 和 self.list
        self.items = self.collect_items_from_list()
        self.save_list_data()
    
    def on_copy_item_content(self, item: DataItem):
        """复制单个数据项的内容到剪贴板"""
        self.on_copy_items_content([item])
    
    def on_copy_items_content(self, items: List[DataItem]):
        """复制多个数据项的内容到剪贴板"""
        try:
            if not items:
                QtWidgets.QMessageBox.information(self, "提示", "没有可复制的内容")
                return
            
            # 提取所有项的内容
            contents = []
            for item in items:
                content = self._extract_item_content(item)
                if content:
                    # 如果是多个项，添加分隔符和标题
                    if len(items) > 1:
                        label = item.display_label()
                        contents.append(f"【{label}】\n{content}\n")
                    else:
                        contents.append(content)
            
            if not contents:
                QtWidgets.QMessageBox.information(self, "提示", "选中的项没有可复制的内容")
                return
            
            # 合并内容
            if len(items) > 1:
                sep = "\n" + "-" * 80 + "\n"
                merged_content = sep.join(contents)
            else:
                merged_content = contents[0]
            
            # 复制到剪贴板
            clipboard = QtWidgets.QApplication.clipboard()
            clipboard.setText(merged_content)
            
            # 记录日志
            if len(items) == 1:
                self.log(f"已复制内容: {items[0].display_label()}")
            else:
                self.log(f"已复制 {len(items)} 个项的内容")
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "错误", f"复制失败：{str(e)}")
    
    def _extract_item_content(self, item: DataItem) -> str:
        """从 DataItem 中提取内容文本"""
        if not item.payload:
            return ""
        
        if item.kind == "manual":
            return item.payload.get("text", "").strip()
        
        elif item.kind == "text_file":
            return item.payload.get("content", "").strip()
        
        elif item.kind.startswith("siyuan"):
            return item.payload.get("content", "").strip()
        
        elif item.kind == "folder":
            # 文件夹：合并所有文件内容
            lines = []
            folder_path = item.payload.get("folder_path", "未知文件夹")
            lines.append(f"【文件夹：{folder_path}】\n")
            for file_info in item.payload.get("files", []):
                file_name = os.path.basename(file_info.get("file_path", "未知文件"))
                file_content = file_info.get("content", "").strip()
                if file_content:
                    lines.append(f"\n【{file_name}】\n{file_content}\n")
            return "\n".join(lines)
        
        else:
            return ""
    
    def save_list_data(self):
        """保存右侧列表数据到配置文件"""
        try:
            # 从列表收集当前所有项
            items_data = []
            for i in range(self.list.count()):
                d = self.list.item(i).data(QtCore.Qt.UserRole)
                if d:
                    items_data.append(d)
            
            # 读取现有配置
            cfg_data = {}
            if os.path.exists(DEFAULT_CONFIG_PATH):
                try:
                    with open(DEFAULT_CONFIG_PATH, "r", encoding="utf-8") as f:
                        cfg_data = json.load(f)
                except Exception:
                    cfg_data = {}
            
            # 更新列表数据
            cfg_data["list_items"] = items_data
            
            # 保存配置
            with open(DEFAULT_CONFIG_PATH, "w", encoding="utf-8") as f:
                json.dump(cfg_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            # 静默失败，不影响用户体验
            pass
    
    def load_list_data(self):
        """从配置文件加载右侧列表数据"""
        try:
            if not os.path.exists(DEFAULT_CONFIG_PATH):
                return
            
            with open(DEFAULT_CONFIG_PATH, "r", encoding="utf-8") as f:
                cfg_data = json.load(f)
            
            items_data = cfg_data.get("list_items", [])
            if not items_data:
                return
            
            # 恢复列表项
            for item_dict in items_data:
                try:
                    item = DataItem.from_dict(item_dict)
                    self.items.append(item)
                    icon = self._icon_for_kind(item.kind)
                    text = item.display_label()
                    lw_item = QtWidgets.QListWidgetItem(icon, text)
                    lw_item.setData(QtCore.Qt.UserRole, item.to_dict())
                    self.list.addItem(lw_item)
                except Exception:
                    # 跳过无效项
                    continue
            
            if items_data:
                self.log(f"已加载 {len(items_data)} 个保存的数据项")
        except Exception:
            # 加载失败时静默处理
            pass
    
    def closeEvent(self, event: QtGui.QCloseEvent):
        """窗口关闭时保存数据"""
        self.save_list_data()
        event.accept()

    def on_settings(self):
        dlg = SettingsDialog(self)
        if dlg.exec_() == QtWidgets.QDialog.Accepted:
            self.siyuan_helper = SiYuanHelper(dlg.cfg)
            self.log("设置已保存")


def main():
    # 高 DPI 缩放设置必须在 QApplication 之前
    QtWidgets.QApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling, True)
    QtWidgets.QApplication.setAttribute(QtCore.Qt.AA_UseHighDpiPixmaps, True)

    app = QtWidgets.QApplication([])

    w = MainWindow()
    w.show()
    app.exec_()


if __name__ == "__main__":
    main()
